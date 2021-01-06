# -*- coding: utf-8 -*-
from openpyxl import *
from docx import Document
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


def strip_all(value):
    temp = value or ''
    return re.sub(r'\s', '', temp)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

class ExcelOp(object):
    def __init__(self, file):
        self.file = file
        self.wb = load_workbook(self.file)
        sheets = self.wb.get_sheet_names()
        self.sheet = sheets[0]
        self.ws = self.wb[self.sheet]

    # 获取表格的总行数和总列数
    def get_row_clo_num(self):
        rows = self.ws.max_row
        columns = self.ws.max_column
        return rows, columns

    # 获取某个单元格的值
    def get_cell_value(self, row, column):
        cell_value = self.ws.cell(row=row, column=column).value
        return cell_value

    # 获取某列的所有值
    def get_col_value(self, column):
        rows = self.ws.max_row
        column_data = []
        for i in range(1, rows + 1):
            cell_value = self.ws.cell(row=i, column=column).value
            column_data.append(cell_value)
        return column_data

    # 获取某行所有值
    def get_row_value(self, row):
        columns = self.ws.max_column
        row_data = []
        for i in range(1, columns + 1):
            cell_value = self.ws.cell(row=row, column=i).value
            row_data.append(cell_value)
        return row_data

    # 设置某个单元格的值
    def set_cell_value(self, row, colunm, cellvalue):
        try:
            self.ws.cell(row=row, column=colunm).value = cellvalue
            self.wb.save(self.file)
        except:
            self.ws.cell(row=row, column=colunm).value = "writefail"
            self.wb.save(self.file)


def to_docx(file, table):
    print('存入以下信息')
    print(table)

    first_row = table[0]
    second_row = table[1]
    data_rows = table[2:]

    # 添加三级标题
    document.add_heading(first_row[3], level=3)
    # 添加段落，表名
    # document.add_paragraph(
    #     first_row[0], style='List Bullet'  # 样式为一个小圆点
    # )
    # 添加表格
    table = document.add_table(rows=1, cols=7)
    # 获取第一行的所有列数
    hdr_cells = table.rows[0].cells
    # 给第一行的各个列添加内容
    hdr_cells[0].text = '名称'
    hdr_cells[1].text = '虚拟'
    hdr_cells[2].text = '类型'
    hdr_cells[3].text = '可为空'
    hdr_cells[4].text = '默认/表达式'
    hdr_cells[5].text = '存储'
    hdr_cells[6].text = '注释'

    for cell in hdr_cells:
        set_cell_border(
            cell,
            top={"sz": 1, "val": "single", "space": "0"},
            bottom={"sz": 1, "val": "single"},
            start={"sz": 1, "val": "single"},
            end={"sz": 1, "val": "single"},
        )

    fields = []
    # 给table表格添加新行，并给各列添加内容
    for one_data in data_rows:
        fields.append(f'{one_data[0]} {one_data[2]}')
        row_cells = table.add_row().cells
        row_cells[0].text = strip_all(one_data[0])
        row_cells[1].text = strip_all(one_data[1])
        row_cells[2].text = strip_all(one_data[2])
        row_cells[3].text = strip_all(one_data[3])
        row_cells[4].text = strip_all(one_data[4])
        row_cells[5].text = strip_all(one_data[5])
        row_cells[6].text = strip_all(one_data[6])

        for cell in row_cells:
            set_cell_border(
                cell,
                top={"sz": 1, "val": "single", "space": "0"},
                bottom={"sz": 1, "val": "single"},
                start={"sz": 1, "val": "single"},
                end={"sz": 1, "val": "single"},
            )

    # 添加段落，表名
    document.add_paragraph(
        '这里放建表语句'
    )
    pass


if __name__ == '__main__':
    excel_op = ExcelOp(file="source.xlsx")

    row_num, col_num = excel_op.get_row_clo_num()

    # 创建word文档对象
    document = Document()

    table = []
    for i in range(1, row_num+1):
        row = excel_op.get_row_value(i)
        if row[0] is None:
            if len(table) != 0:
                to_docx(document, table)
            table = []
        else:
            table.append(row)
            pass
        pass

    if len(table) != 0:
        to_docx(document, table)

    # 保存world文档
    document.save('demo.docx')

